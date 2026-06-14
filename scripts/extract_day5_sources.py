from pathlib import Path
from docx import Document

IMPLEMENTATION_DIR = Path(__file__).resolve().parents[1]
PROJECT_DIR = IMPLEMENTATION_DIR.parent

SOURCE_FILES = [
    "04_System_Architecture_and_Experimental_Boundary.docx",
    "05_Failure_Injection_and_Experiment_Matrix.docx",
    "06_Experiment_Data_Schema.docx",
    "07_Technical_Feasibility_and_Implementation_Stack.docx",
]

OUTPUT_FILE = (
    IMPLEMENTATION_DIR
    / "research_inputs"
    / "day5_methodology_sources.md"
)


def clean_text(value: str) -> str:
    return " ".join(value.split())


def escape_table_cell(value: str) -> str:
    return clean_text(value).replace("|", r"\|")


def extract_document(path: Path) -> list[str]:
    document = Document(path)
    lines: list[str] = [f"# {path.stem}", ""]

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue

        style = paragraph.style.name.lower()

        if style.startswith("title"):
            lines.extend([f"# {text}", ""])
        elif style.startswith("heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 2
            level = min(max(level + 1, 2), 6)
            lines.extend([f"{'#' * level} {text}", ""])
        else:
            lines.extend([text, ""])

    for table_number, table in enumerate(document.tables, start=1):
        lines.extend([f"## Extracted Table {table_number}", ""])

        rows = [
            [escape_table_cell(cell.text) for cell in row.cells]
            for row in table.rows
        ]

        if not rows:
            continue

        column_count = max(len(row) for row in rows)
        rows = [
            row + [""] * (column_count - len(row))
            for row in rows
        ]

        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * column_count) + " |")

        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")

        lines.append("")

    return lines


def main() -> None:
    output_lines = [
        "# Day 5 Methodology Source Extraction",
        "",
        "Documents 04–07 extracted for methodology and experimental-results drafting.",
        "",
    ]

    for filename in SOURCE_FILES:
        source_path = PROJECT_DIR / filename

        if not source_path.exists():
            raise FileNotFoundError(f"Missing source document: {source_path}")

        output_lines.extend(extract_document(source_path))
        output_lines.extend(["---", ""])

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_FILE.write_text(
        "\n".join(output_lines),
        encoding="utf-8",
    )

    print(f"Created: {OUTPUT_FILE}")
    print(f"Extracted documents: {len(SOURCE_FILES)}")


if __name__ == "__main__":
    main()
